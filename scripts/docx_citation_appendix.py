"""DOCX adapter for citation-auditor aggregated JSON.

`citation_auditor` is markdown-native. This project ships formal legal
opinions primarily as DOCX, so DOCX renderers can use this sibling adapter to
consume the aggregated verdict JSON without modifying the vendored auditor.

Public hooks:
- load_aggregated(path) -> dict
- inject_unverified_tags(body_md, aggregated) -> str
- append_citation_audit_log(doc, aggregated) -> None
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor

FONT_CJK = "맑은 고딕"
FONT_LATIN = "Times New Roman"

_TAG_FOR_LABEL = {
    "contradicted": "[Unverified]",
    "unknown": "[Partially Unverified]",
}

_APPENDIX_HEADING = "부록: 검증 로그 (Citation Audit Log)"
_APPENDIX_DISCLAIMER = (
    "본 산출물에 포함된 사실·인용 주장에 대한 자동 인용 감사 결과입니다. "
    "자동 감사는 사람의 검토를 대체하지 않으며, Contradicted / Unknown 항목은 "
    "최종 검토 시 확인이 필요합니다."
)
_APPENDIX_HEADERS = ["#", "클레임 (Claim)", "판정 (Verdict)", "Verifier", "근거 (Evidence)"]
_COL_WIDTHS_CM = (1.0, 7.0, 2.5, 2.5, 5.0)


def load_aggregated(path: str | Path) -> dict[str, Any]:
    """Load a citation-auditor aggregate output JSON file."""
    return json.loads(Path(path).read_text(encoding="utf-8"))


def inject_unverified_tags(body_md: str, aggregated: dict[str, Any]) -> str:
    """Insert audit tags after failing claims in markdown body text.

    The aggregate stage converts chunk-relative spans to document-relative
    offsets. Invalid spans are ignored rather than guessed.
    """
    insertions: list[tuple[int, str]] = []
    for item in aggregated.get("aggregated", []):
        verdict = item.get("verdict") or {}
        tag = _TAG_FOR_LABEL.get(verdict.get("label"))
        if tag is None:
            continue
        claim = verdict.get("claim") or {}
        span = claim.get("sentence_span") or {}
        end = span.get("end")
        if isinstance(end, int) and 0 <= end <= len(body_md):
            insertions.append((end, f" {tag}"))

    rendered = body_md
    for offset, tag in sorted(insertions, key=lambda item: item[0], reverse=True):
        rendered = rendered[:offset] + tag + rendered[offset:]
    return rendered


def append_citation_audit_log(doc: DocxDocument, aggregated: dict[str, Any]) -> None:
    """Append a bilingual Citation Audit Log appendix to a DOCX document."""
    doc.add_page_break()
    heading = doc.add_heading(_APPENDIX_HEADING, level=2)
    for run in heading.runs:
        _set_run_font(run, size=Pt(13), bold=True, color=RGBColor(0x00, 0x2B, 0x5C))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(_APPENDIX_DISCLAIMER)
    _set_run_font(run, size=Pt(9.5), color=RGBColor(0x59, 0x59, 0x59))

    verdicts = aggregated.get("aggregated") or []
    if not verdicts:
        p = doc.add_paragraph()
        run = p.add_run("[Citation Audit Skipped — no claims extracted or verifiers unavailable]")
        _set_run_font(run, size=Pt(10), bold=True, color=RGBColor(0xC0, 0x00, 0x00))
        return

    table = doc.add_table(rows=1, cols=len(_APPENDIX_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, width in enumerate(_COL_WIDTHS_CM):
        table.columns[index].width = Cm(width)

    header_cells = table.rows[0].cells
    for index, header in enumerate(_APPENDIX_HEADERS):
        cell = header_cells[index]
        _set_cell_shading(cell, "002B5C")
        _set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        _set_run_font(run, size=Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for row_index, item in enumerate(verdicts, start=1):
        verdict = item.get("verdict") or {}
        claim = verdict.get("claim") or {}
        row = table.add_row()
        values = [
            str(row_index),
            _truncate(claim.get("text", ""), 140),
            _format_label(verdict.get("label", "unknown")),
            verdict.get("verifier_name", "-"),
            _format_evidence(verdict.get("evidence") or []),
        ]
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            _set_cell_margins(cell)
            if row_index % 2 == 0:
                _set_cell_shading(cell, "F0F4F8")
            paragraph = cell.paragraphs[0]
            if col_index in (0, 2, 3):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            _set_run_font(run, size=Pt(8.5))


def _set_run_font(run, *, size: Pt, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_LATIN
    run.font.size = size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(
        parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear" '
            'w:color="auto"/>'
        )
    )


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 80, right: int = 80) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(
        parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            '</w:tcMar>'
        )
    )


def _truncate(text: str, max_len: int) -> str:
    collapsed = " ".join(str(text).split())
    if len(collapsed) <= max_len:
        return collapsed
    return collapsed[: max_len - 1] + "…"


def _format_label(label: str) -> str:
    normalized = str(label or "unknown")
    return normalized.upper()


def _format_evidence(evidence: list[dict[str, Any]]) -> str:
    if not evidence:
        return "-"
    parts: list[str] = []
    for item in evidence:
        url = str(item.get("url") or "").strip()
        title = str(item.get("title") or "").strip()
        if title and url and title != url:
            parts.append(f"{title} ({url})")
        elif url:
            parts.append(url)
        elif title:
            parts.append(title)
    return _truncate("; ".join(parts), 120) if parts else "-"
