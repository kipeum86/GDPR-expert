#!/usr/bin/env python3
"""
Generate Korean legal analysis memo DOCX following ko-legal-opinion-style-guide.md.

Style: Korean professional-format legal analysis memo.
Font: 맑은 고딕 (CJK) + Times New Roman (Latin), 11pt.
Paper: A4, margins 2.54cm.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cjk_font(run, font_name_cjk="맑은 고딕", font_name_latin="Times New Roman", size=Pt(11)):
    """Set CJK + Latin fonts on a run (style guide §11.3)."""
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name_cjk)
    rFonts.set(qn("w:ascii"), font_name_latin)
    rFonts.set(qn("w:hAnsi"), font_name_latin)


def add_run(paragraph, text, bold=False, size=Pt(11), color=None):
    """Add a run with CJK font settings."""
    run = paragraph.add_run(text)
    set_cjk_font(run, size=size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(doc, text="", bold=False, size=Pt(11), alignment=None,
                  space_before=None, space_after=None, line_spacing=None,
                  left_indent=None):
    """Add a paragraph with standard formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if text:
        add_run(p, text, bold=bold, size=size)
    return p


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right with val, sz, color."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        elem = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(elem)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """Set cell internal margins in twips."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_statute_block(doc, lines):
    """Add a statute/law citation block in a bordered single-cell table (§4.1)."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=80, bottom=80, left=150, right=150)
    for edge in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, **{edge: {"val": "single", "sz": "4", "color": "000000"}})

    # Clear default paragraph
    cell.paragraphs[0].clear()

    for i, line_data in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

        if isinstance(line_data, dict):
            text = line_data.get("text", "")
            bold = line_data.get("bold", False)
            indent = line_data.get("indent", 0)
            if indent > 0:
                p.paragraph_format.left_indent = Cm(indent * 0.5)
            add_run(p, text, bold=bold, size=Pt(10))
        else:
            add_run(p, line_data, size=Pt(10))

    doc.add_paragraph()  # spacing after block


def add_analysis_table(doc, rows_data):
    """Add a formatted analysis table (e.g., balancing test)."""
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            p = cell.paragraphs[0]
            for edge in ['top', 'bottom', 'left', 'right']:
                set_cell_border(cell, **{edge: {"val": "single", "sz": "4", "color": "999999"}})
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
            if ri == 0:
                add_run(p, cell_text, bold=True, size=Pt(10))
            else:
                add_run(p, cell_text, size=Pt(10))

    doc.add_paragraph()


# ── Document Setup ───────────────────────────────────────────────────────

def setup_document():
    """Create document with A4 paper, 2.54cm margins, base font styles."""
    doc = Document()

    # Page setup: A4, 2.54cm margins (§11.1)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Base style (§11.2)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # Paragraph format (§11.4)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    return doc


# ── Content ──────────────────────────────────────────────────────────────

def build_document():
    doc = setup_document()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # §2.1 MEMORANDUM header
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    p = add_paragraph(doc, "법률 분석 메모", bold=True, size=Pt(18),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_after=Pt(4))

    # §2.2 Date
    add_paragraph(doc, "2026. 3. 26.", size=Pt(11),
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_after=Pt(12))

    # §2.3 Information block (bordered table)
    info_table = doc.add_table(rows=3, cols=2)
    info_data = [
        ("요청 맥락", "[Client] — B2B SaaS Procurement Platform"),
        ("생성 주체", "Jinju Legal Orchestrator — EU 데이터 보호 스페셜리스트 Kim De Bruyne (김덕배)"),
        ("제 목", "GDPR상 AI 기반 지출 분석 기능을 위한 고객 조달 데이터의 LLM 미세조정 적법성 검토"),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_l = info_table.cell(i, 0)
        cell_v = info_table.cell(i, 1)
        for edge in ['top', 'bottom', 'left', 'right']:
            set_cell_border(cell_l, **{edge: {"val": "single", "sz": "4", "color": "000000"}})
            set_cell_border(cell_v, **{edge: {"val": "single", "sz": "4", "color": "000000"}})
        set_cell_margins(cell_l, top=40, bottom=40, left=100, right=60)
        set_cell_margins(cell_v, top=40, bottom=40, left=60, right=100)
        # Label
        p_l = cell_l.paragraphs[0]
        add_run(p_l, label, bold=True, size=Pt(11))
        add_run(p_l, " :", bold=True, size=Pt(11))
        # Value
        p_v = cell_v.paragraphs[0]
        add_run(p_v, value, size=Pt(11))

    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(13.5)

    doc.add_paragraph()  # spacing

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # §3.3 분석 메모의 한계 (Limitations Disclaimer)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc,
        "아래 분석은 귀사가 제공한 자료 및 정보만을 전제로 귀사가 문의한 사항에 "
        "국한된 법률검토임을 말씀드립니다. 이 사건에 제공된 자료 및 정보 이외에 "
        "다른 특별한 사정이 있는 경우 그 법률적 판단이 달라질 수 있습니다.",
        size=Pt(11), space_after=Pt(12))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 1. 배경 사실 관계 (§3.1)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "1. 배경 사실 관계", bold=True, size=Pt(13), space_before=Pt(18), space_after=Pt(8))

    add_paragraph(doc,
        "1. 귀사는 B2B SaaS 플랫폼을 운영하며, 기업 고객을 대리하여 조달 데이터"
        "(인보이스, 구매주문서, 벤더 계약 등)를 처리하고 있습니다. 귀사는 「General "
        "Data Protection Regulation」(이하 \"GDPR\") 제28조에 따른 개인정보 수탁처리자"
        "(processor)로서 활동하고 있습니다.",
        left_indent=Cm(0.5))

    add_paragraph(doc,
        "2. 해당 데이터에는 고객사 직원의 업무 연락처(성명, 이메일, 직함, 전화번호)가 "
        "포함되어 있습니다.",
        left_indent=Cm(0.5))

    add_paragraph(doc,
        "3. 귀사의 CTO는 위 데이터를 사용하여 \"AI 기반 지출 분석\" 기능 — 지출 패턴 "
        "예측 및 비용 최적화 제안 — 을 위한 대규모 언어모델(이하 \"LLM\")을 미세조정"
        "(fine-tuning)할 것을 제안하고 있습니다.",
        left_indent=Cm(0.5))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 2. 질의의 요지 (§3.2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "2. 질의의 요지", bold=True, size=Pt(13), space_before=Pt(18), space_after=Pt(8))

    add_paragraph(doc, "귀사는 다음과 같은 사항을 질의하셨습니다.", left_indent=Cm(0.5))

    add_paragraph(doc,
        "1. 현재의 수탁처리자(processor) 자격으로 고객 조달 데이터를 사용하여 LLM을 "
        "미세조정하는 것이 GDPR상 적법한지 여부",
        left_indent=Cm(1.0))
    add_paragraph(doc,
        "2. GDPR 제6조 제1항 (f)호의 정당한 이익(legitimate interest)을 법적 근거로 "
        "원용할 수 있는지 여부",
        left_indent=Cm(1.0))
    add_paragraph(doc,
        "3. 적법한 AI 모델 학습을 위하여 필요한 법률관계 재구성 및 완화 조치",
        left_indent=Cm(1.0))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 3. 검토 결과 (§3.4)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "3. 검토 결과", bold=True, size=Pt(13), space_before=Pt(18), space_after=Pt(8))

    # ── 가. 선결 쟁점: Controller vs. Processor ──
    add_paragraph(doc,
        "가. 선결 쟁점: 개인정보처리자(Controller) 및 수탁처리자(Processor) 지위 검토",
        bold=True, size=Pt(12), space_before=Pt(14), space_after=Pt(6),
        left_indent=Cm(0.5))

    add_paragraph(doc, "(1) 관련 법령", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    # Statute block: Art. 28, 29, 28(10)
    add_statute_block(doc, [
        {"text": "「General Data Protection Regulation」 (Regulation (EU) 2016/679)", "bold": True},
        {"text": ""},
        {"text": "Article 28(3)(a) — Processor", "bold": True},
        {"text": "\"[The processor shall] processes the personal data only on documented instructions "
                 "from the controller, including with regard to transfers of personal data to a third "
                 "country or an international organisation\"", "indent": 1},
        {"text": ""},
        {"text": "Article 29 — Processing under the authority of the controller or processor", "bold": True},
        {"text": "\"The processor and any person acting under the authority of the controller or of the "
                 "processor, who has access to personal data, shall not process those data except on "
                 "instructions from the controller, unless required to do so by Union or Member State law.\"",
         "indent": 1},
        {"text": ""},
        {"text": "Article 28(10)", "bold": True},
        {"text": "\"[I]f a processor infringes this Regulation by determining the purposes and means of "
                 "processing, the processor shall be considered to be a controller in respect of that "
                 "processing.\"", "indent": 1},
    ])

    add_paragraph(doc, "(2) 전문(Recital) 맥락", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_statute_block(doc, [
        {"text": "「GDPR」 Recital 81", "bold": True},
        {"text": "\"The carrying-out of processing by a processor should be governed by a contract or "
                 "other legal act under Union or Member State law, binding the processor to the controller, "
                 "setting out the subject-matter and duration of the processing, the nature and purposes of "
                 "the processing, the type of personal data and categories of data subjects, taking into "
                 "account the specific tasks and responsibilities of the processor in the context of the "
                 "processing to be carried out and the risk to the rights and freedoms of the data subject.\"",
         "indent": 1},
    ])

    add_paragraph(doc,
        "위 Recital 81은 수탁처리자의 역할이 개인정보처리자와의 계약에 의해 한정됨을 확인합니다. "
        "즉, 수탁처리자는 계약에서 정의된 개인정보처리자의 목적을 위해 데이터를 처리하는 것이지, "
        "자체 목적을 위해 처리하는 것이 아닙니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(3) EDPB 가이드라인", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "EDPB Guidelines 07/2020 (개인정보처리자 및 수탁처리자 개념에 관한 가이드라인)은 "
        "다음의 기준을 제시하고 있습니다:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "1) 개인정보처리자 개념은 기능적 개념(functional concept)으로, 형식적 지정이 아닌 "
        "실제 역할에 따라 책임을 배분하는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "2) 개인정보처리자는 처리의 목적과 수단 — \"왜(why)\"와 \"어떻게(how)\" — 를 "
        "결정하는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "3) 수탁처리자가 지시를 넘어 자체적으로 목적과 수단을 결정하기 시작하면, 해당 "
        "처리에 대해 개인정보처리자로 간주되는 점",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "등을 고려할 때, 수탁처리자의 역할 범위는 엄격히 제한되어야 합니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(4) 검토 결과", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "고객 데이터를 사용하여 자사의 LLM을 미세조정하는 것은 처리에 대한 새로운 목적을 "
        "결정하는 것에 해당합니다. 원래의 목적(조달 플랫폼 서비스 제공)은 개인정보처리자인 "
        "고객사가 정한 것인바, 자사 제품 개발을 위한 AI 모델 학습은 귀사가 결정하는 완전히 "
        "별개의 목적에 해당합니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "따라서, GDPR 제28조 제10항에 따라, 귀사는 AI 학습 처리에 대해 개인정보처리자로 "
        "간주될 것으로 판단됩니다. 이는 이론적 위험이 아니라 GDPR이 규정한 직접적 "
        "법률효과입니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "한편, 이탈리아 개인정보보호 감독기관(Garante)이 OpenAI에 대하여 적절한 법적 근거 "
        "없이 모델 학습을 위해 개인정보를 처리한 것을 이유로 1,500만 유로의 과징금을 부과한 "
        "사례(Garante v OpenAI, 2024. 12. 20. 결정)는, 감독기관이 AI 기업에 대해 적극적으로 "
        "집행하고 있음을 보여줍니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "법률관계 재구성 없이 고객 데이터를 AI 학습에 사용할 경우 다음과 같은 결과가 "
        "발생할 수 있습니다:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "1) 처리 계약 위반 — 제28조 제3항 (a)호 위반(지시 범위를 초과한 처리)으로 인한 "
        "계약상 책임 및 해지 사유 발생 가능성이 있는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "2) 행정 과징금 — 제83조 제4항 (a)호에 따라 최대 1,000만 유로 또는 전 세계 "
        "연간 매출의 2%에 해당하는 과징금이 부과될 수 있는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "3) 개인정보처리자의 전체 의무 부담 — AI 학습의 간주 개인정보처리자로서 모든 "
        "GDPR 개인정보처리자 의무(법적 근거, 투명성, 정보주체 권리, DPIA 등)를 부담하게 "
        "되는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "4) 유효한 법적 근거 부재 — 귀사는 고객의 법적 근거(제6조 제1항 (b)호 — 계약상 "
        "필요성)에 따라 수탁처리자로서 데이터를 수집하였으므로, 자체 목적으로 사용하려면 "
        "독자적인 법적 근거가 필요하나 현재 이를 갖추고 있지 아니한 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "5) 민사 책임 — 제82조에 따라 정보주체가 물질적 및 비물질적 손해에 대한 배상을 "
        "청구할 수 있는 점",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "등을 고려할 때, 위험 수준은 높다고 판단됩니다.",
        left_indent=Cm(1.0))

    # ── 나. 목적 제한 원칙 ──
    add_paragraph(doc,
        "나. 목적 제한 원칙 (Art. 5(1)(b)) 검토",
        bold=True, size=Pt(12), space_before=Pt(14), space_after=Pt(6),
        left_indent=Cm(0.5))

    add_paragraph(doc, "(1) 관련 법령", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_statute_block(doc, [
        {"text": "「GDPR」", "bold": True},
        {"text": ""},
        {"text": "Article 5(1)(b) — Purpose limitation", "bold": True},
        {"text": "\"[Personal data shall be] collected for specified, explicit and legitimate purposes and "
                 "not further processed in a manner that is incompatible with those purposes\"", "indent": 1},
        {"text": ""},
        {"text": "Recital 50", "bold": True},
        {"text": "\"The processing of personal data for purposes other than those for which the personal "
                 "data were initially collected should be allowed only where the processing is compatible "
                 "with the purposes for which the personal data were initially collected.\"", "indent": 1},
        {"text": ""},
        {"text": "Recital 39 (excerpt)", "bold": True},
        {"text": "\"[T]he specific purposes for which personal data are processed should be explicit and "
                 "legitimate and determined at the time of the collection of the personal data. [...] "
                 "Personal data should be processed only if the purpose of the processing could not "
                 "reasonably be fulfilled by other means.\"", "indent": 1},
    ])

    add_paragraph(doc,
        "위 Recital 39의 마지막 문장 — \"처리 목적이 다른 수단으로 합리적으로 달성될 수 없는 "
        "경우에만 개인정보를 처리해야 한다\" — 는 AI 모델 학습에 특히 중요합니다. 이는 목적 제한과 "
        "필요성 평가 간의 직접적 연결을 생성하여, 개인정보 없이(예: 익명화 또는 합성 데이터를 "
        "통해) 동일한 목적을 달성할 수 있는지 개인정보처리자가 검토할 것을 요구하기 때문입니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(2) Art. 6(4)에 따른 호환성 평가", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_statute_block(doc, [
        {"text": "「GDPR」 Article 6(4)", "bold": True},
        {"text": "원래 수집 목적과 다른 목적으로의 처리 시 고려해야 할 요소:", "indent": 1},
        {"text": "(a) 수집 목적과 추가 처리 목적 간의 연관성", "indent": 1},
        {"text": "(b) 특히 정보주체와 개인정보처리자 간의 관계를 고려한 수집 맥락", "indent": 1},
        {"text": "(c) 개인정보의 성격", "indent": 1},
        {"text": "(d) 추가 처리가 정보주체에게 미칠 수 있는 결과", "indent": 1},
        {"text": "(e) 암호화 또는 가명처리를 포함한 적절한 보호조치의 존재", "indent": 1},
    ])

    add_paragraph(doc, "(3) 검토 결과", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "Art. 6(4)의 호환성 테스트를 본건에 적용하여 보건대:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "1) 목적 간 연관성(a): 약함 — 원래 목적은 조달 SaaS 플랫폼 제공(운영)이고, "
        "LLM 미세조정은 제품 개발 활동인 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "2) 수집 맥락(b): 불리 — 데이터는 B2B 맥락에서 수집되었으며, 정보주체는 자신의 "
        "데이터가 AI 모델 학습에 사용될 것을 합리적으로 기대하지 아니하는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "3) 데이터 성격(c): 중간 위험 — 업무 연락처는 민감정보(Art. 9)는 아니나, 조달 "
        "데이터는 상업적으로 민감한 정보를 포함할 수 있는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "4) 결과(d): 중간 — 데이터가 모델 가중치에 내재될 수 있어 추출 또는 재출력 "
        "위험이 존재하는 점,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "5) 보호조치(e): 유리(구현 시) — 가명처리, 차등 프라이버시 등의 기술적 조치로 "
        "위험을 완화할 수 있는 점",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "등에 비추어 볼 때, AI 모델 학습은 조달 플랫폼 제공이라는 원래 목적과 명백히 "
        "호환되지 아니하며, 별도의 법적 근거를 요하는 독립적 목적으로 판단됩니다.",
        left_indent=Cm(1.0))

    # ── 다. 정당한 이익 균형 테스트 ──
    add_paragraph(doc,
        "다. 정당한 이익 균형 테스트 (Art. 6(1)(f)) 검토",
        bold=True, size=Pt(12), space_before=Pt(14), space_after=Pt(6),
        left_indent=Cm(0.5))

    add_paragraph(doc,
        "이하의 분석은 상기 가.항의 개인정보처리자/수탁처리자 쟁점이 적절히 재구성되어, "
        "귀사가 AI 학습 목적에 대해 개인정보처리자(또는 공동 개인정보처리자)로 활동하는 "
        "경우를 전제로 합니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(1) 관련 법령 및 EDPB 가이드라인", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_statute_block(doc, [
        {"text": "「GDPR」 Article 6(1)(f)", "bold": True},
        {"text": "\"[P]rocessing is necessary for the purposes of the legitimate interests pursued by the "
                 "controller or by a third party, except where such interests are overridden by the interests "
                 "or fundamental rights and freedoms of the data subject which require protection of personal "
                 "data, in particular where the data subject is a child.\"", "indent": 1},
        {"text": ""},
        {"text": "Recital 47 (excerpt)", "bold": True},
        {"text": "\"At any rate the existence of a legitimate interest would need careful assessment including "
                 "whether a data subject can reasonably expect at the time and in the context of the collection "
                 "of the personal data that processing for that purpose may take place.\"", "indent": 1},
    ])

    add_paragraph(doc,
        "EDPB Guidelines 1/2024 (정당한 이익에 관한 가이드라인, 2024. 10. 8.)는 다음의 "
        "3단계 테스트를 제시하고 있습니다:",
        left_indent=Cm(1.0))
    add_paragraph(doc,
        "① 개인정보처리자 또는 제3자의 정당한 이익 추구",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "② 해당 정당한 이익 목적을 위한 처리의 필요성",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "③ 정보주체의 이익 또는 기본권과 자유가 우선하지 않을 것",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "나아가, EDPB Opinion 28/2024 (2024. 12. 17.)는 위 기본 틀을 AI 모델 개발 및 배포에 "
        "구체적으로 적용하고 있습니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(2) 3단계 분석", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc, "가) 제1단계: 정당한 이익 식별", bold=True, left_indent=Cm(1.5), space_before=Pt(6))

    add_paragraph(doc,
        "이익은 ① 적법하고, ② 명확하고 정밀하게 표현되었으며, ③ 현실적이고 현재의"
        "(추측적이지 않은) 것이어야 합니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "EDPB Opinion 28/2024 (para. 69)는 AI 맥락에서의 정당한 이익의 예시로 (i) 사용자를 "
        "지원하는 대화형 에이전트 서비스 개발, (ii) 사기성 콘텐츠 또는 행위를 탐지하는 AI 시스템 "
        "개발, (iii) 정보 시스템의 위협 탐지 개선 등을 제시하고 있습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "귀사의 이익 — \"고객을 위한 지출 패턴 예측 및 비용 최적화 제안을 위한 AI 기반 지출 "
        "분석 개발\" — 은 적법하고, 명확히 표현되어 있으며, 현실적이고 현재의 이익에 해당하는바, "
        "정당한 이익으로 인정될 가능성이 높을 것으로 판단됩니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc, "나) 제2단계: 처리의 필요성", bold=True, left_indent=Cm(1.5), space_before=Pt(6))

    add_paragraph(doc,
        "EDPB Opinion 28/2024 (paras. 70-75)는 처리가 정당한 이익 추구를 허용하는지, 그리고 "
        "덜 침해적인 방법이 없는지를 검토할 것을 요구합니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "이 점에 관하여 EDPB Opinion 28/2024 (para. 73)는 다음과 같이 판시하고 있습니다:",
        left_indent=Cm(1.5))

    add_statute_block(doc, [
        {"text": "EDPB Opinion 28/2024 (para. 73)", "bold": True},
        {"text": "\"If the pursuit of the purpose is also possible through an AI model that does not entail "
                 "processing of personal data, then processing personal data should be considered as not "
                 "necessary.\"", "indent": 1},
    ])

    add_paragraph(doc,
        "본 사안에서 지출 분석(지출 패턴 예측, 비용 최적화 제안)의 핵심 가치는 거래 데이터"
        "(금액, 카테고리, 시기, 벤더 유형)에 있으며, 개인 식별정보에 있지 아니합니다. 가명처리 "
        "또는 익명화된 조달 데이터로 모델을 효과적으로 학습할 수 있을 가능성이 높은바, "
        "필요성이 가장 취약한 부분으로 판단됩니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc, "다) 제3단계: 균형 테스트", bold=True, left_indent=Cm(1.5), space_before=Pt(6))

    add_paragraph(doc,
        "필요성이 확립된 경우를 가정하여 균형 테스트를 수행하면 다음과 같습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "EDPB Opinion 28/2024 (paras. 77-80)는 관련 정보주체 이익으로서 ① 개인정보에 대한 "
        "자기결정 및 통제, ② AI 모델이 직업 활동에 영향을 미치는 경우의 경제적 이익, "
        "③ EU 기본권헌장 제7조(사생활)와 제8조(개인정보보호)에 따른 기본권, ④ 학습된 "
        "모델로부터의 개인정보 추출 또는 재출력 위험 등을 식별하고 있습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "본 사안에서 정보주체는 고객사의 직원으로서, 소비자나 취약한 개인이 아닌 "
        "직업적(B2B) 맥락의 개인인바, 이는 개인정보처리자의 이익에 다소 유리하게 "
        "작용합니다. 그러나, Recital 47에 비추어 보건대 정보주체가 자신의 조달 데이터가 "
        "AI 모델 학습에 사용될 것을 합리적으로 기대하지 않을 것이라는 점은 정당한 이익 "
        "원용에 상당히 불리하게 작용합니다.",
        left_indent=Cm(1.5))

    # Balancing table
    add_analysis_table(doc, [
        ["요소", "비중", "방향"],
        ["AI 혁신에 대한 상업적 이익", "중간", "개인정보처리자에 유리"],
        ["B2B 직업적 맥락 (소비자 아님)", "중간", "개인정보처리자에 유리"],
        ["고객에 대한 혜택 (비용 최적화)", "중간", "개인정보처리자에 유리"],
        ["합리적 기대의 부재", "강함", "정보주체에 유리"],
        ["정보주체와의 간접적 관계", "중간", "정보주체에 유리"],
        ["모델로부터의 데이터 추출 위험", "중간", "정보주체에 유리"],
        ["상업적으로 민감한 데이터", "중간", "정보주체에 유리"],
        ["이용 가능한 완화 조치", "중간", "개인정보처리자에 유리 (구현 시)"],
    ])

    add_paragraph(doc,
        "따라서, 강력한 완화 조치 없이는 균형 테스트가 개인정보처리자에게 불리하게 "
        "기울 것으로 판단됩니다. 포괄적 완화(가명처리, 차등 프라이버시, 투명성, 옵트아웃) "
        "시 균형이 이동할 수 있으나, 철저한 문서화와 DPIA가 필요한 근접한 판단(close call)"
        "으로 남을 것으로 사료됩니다.",
        left_indent=Cm(1.5))

    # ── 라. EDPB의 AI 및 개인정보에 관한 입장 ──
    add_paragraph(doc,
        "라. EDPB의 AI 및 개인정보에 관한 입장",
        bold=True, size=Pt(12), space_before=Pt(14), space_after=Pt(6),
        left_indent=Cm(0.5))

    add_paragraph(doc, "(1) EDPB Opinion 28/2024 (2024. 12. 17.)", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "위 분석 메모는 AI 모델 학습과 개인정보에 관한 EDPB의 확정적 입장으로서, 다음과 같은 "
        "핵심 판단을 포함하고 있습니다:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "1) 개인정보로 학습된 AI 모델이 자동으로 익명이 되는 것은 아니며, 모델의 "
        "익명성 여부는 구체적 기준에 따라 사안별로 평가되어야 하는 점 (para. 34),",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "2) AI 모델이 학습 데이터로부터 식별 가능한 자연인에 관한 정보를 출력하도록 의도적으로 "
        "설계되지 않은 경우에도, 개인정보를 포함한 학습 데이터 세트의 정보가 모델의 파라미터에 "
        "\"흡수(absorbed)\"되어 잔존할 수 있는 점 (para. 31),",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "3) 정당한 이익은 가능하나, 처리되는 개인정보의 양과 비례성, 익명화 또는 합성 대안의 "
        "존재 여부, 정보주체의 합리적 기대, 프라이버시 보존 기술, 제21조에 따른 이의제기권 등에 "
        "특별히 주의하여 엄격한 평가가 필요한 점",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "등이 주요 내용에 해당합니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "나아가, 위법한 처리의 결과에 관하여 EDPB는 3가지 시나리오를 제시하고 있는바:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "① 동일 개인정보처리자가 데이터를 보유하는 경우: 후속 처리의 적법성은 사안별로 평가,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "② 다른 개인정보처리자가 배포하는 경우: 배포하는 개인정보처리자는 개발이 적법했는지 "
        "실사(due diligence)를 수행해야 함,",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "③ 위법한 학습 후 모델이 익명화된 경우: 모델이 진정으로 익명인 경우, GDPR이 후속 "
        "운영에 적용되지 아니함",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "이라고 판시하고 있습니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(2) 집행 사례: Garante v OpenAI (2024. 12. 20.)", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "이탈리아 개인정보보호 감독기관(Garante)은 OpenAI에 대하여 다음과 같은 위반을 이유로 "
        "1,500만 유로의 과징금을 부과하였습니다:",
        left_indent=Cm(1.0))

    add_paragraph(doc,
        "1) 유효한 법적 근거 없이 ChatGPT를 개인정보로 학습한 점 (Art. 6 위반),",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "2) AI 학습을 위한 데이터 사용에 대한 투명성이 미흡한 점 (Art. 13 위반),",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "3) 13세 미만 사용자에 대한 연령 확인이 부재한 점 (Art. 8 위반),",
        left_indent=Cm(1.5))
    add_paragraph(doc,
        "4) 데이터 유출 통지를 미이행한 점 (Art. 33 위반)",
        left_indent=Cm(1.5))

    # ── 마. 디지털 옴니버스 패키지 ──
    add_paragraph(doc,
        "마. 디지털 옴니버스 패키지 (COM(2025) 837)의 영향",
        bold=True, size=Pt(12), space_before=Pt(14), space_after=Pt(6),
        left_indent=Cm(0.5))

    add_paragraph(doc, "(1) 개요", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "디지털 법률 옴니버스(COM(2025) 837, 2025. 11. 19. 공개)는 AI 개발과 관련된 중요한 "
        "GDPR 개정안을 제안하고 있습니다. 다만, 이는 입법 제안으로서 아직 법률이 아니며, "
        "입법 과정에서 상당한 변경이 있을 수 있습니다.",
        left_indent=Cm(1.0))

    add_paragraph(doc, "(2) 주요 제안 변경사항", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_paragraph(doc,
        "가) AI 개발을 명시적 정당한 이익으로 인정: 제안 Recital 30은 AI 시스템 및 기반 모델의 "
        "개발과 사용이 GDPR 제6조의 정당한 이익에 해당할 수 있음을 확인하고 있습니다. 다만, "
        "이는 자동적 법적 근거를 생성하는 것이 아니라 AI 개발이 정당한 이익이 될 수 있음을 "
        "확인하는 것에 불과하며, 전체 3단계 테스트는 여전히 적용됩니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "나) AI 처리에 대한 균형 요소: 제안 Recital 31은 이익이 정보주체와 사회에 유익한지, "
        "정보주체의 합리적 기대, 향상된 투명성 제공, 무조건적 이의제기권, AI 학습을 위한 "
        "최신 프라이버시 보존 기술 등의 보호조치를 고려할 것을 제시하고 있습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "다) 특수범주 데이터에 대한 AI 예외: 제안 Art. 9(2)(k)는 AI 시스템 또는 AI 모델의 "
        "개발 및 운영 맥락에서의 처리에 대한 예외를 규정하되, 새로운 Art. 9(5)에 따른 조건"
        "(기술적·조직적 조치, 삭제 의무, 출력 보호 등)을 요구하고 있습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc,
        "라) 개인정보 정의 — 상대적 접근법: 제안된 개정안은 가명처리된 데이터에 대한 상대적/"
        "주관적 접근법을 채택하여, 해당 주체가 자연인을 식별할 \"합리적으로 사용될 가능성이 "
        "있는 수단\"을 갖추지 않은 경우 해당 정보를 개인정보로 보지 아니합니다. 이는 적절히 "
        "가명처리된 데이터를 처리하는 AI 개발자에게 유리할 수 있습니다.",
        left_indent=Cm(1.5))

    add_paragraph(doc, "(3) 현행법 대비 영향 평가", bold=True, left_indent=Cm(1.0), space_before=Pt(8))

    add_analysis_table(doc, [
        ["쟁점", "현행 GDPR", "COM(2025) 837 채택 시"],
        ["AI 학습의 정당한 이익", "명시적 열거 없음; 논증 가능", "Recital에서 명시적 확인"],
        ["균형 테스트 요건", "전체 3단계 테스트", "동일, AI 특화 지침 추가"],
        ["연구 목적의 목적 제한", "좁은 \"과학적 연구\"", "넓은 정의, 상업적 목적 포함 가능"],
        ["AI에서의 특수범주 데이터", "매우 제한적 예외", "잔여 처리를 위한 새 Art. 9(2)(k) 예외"],
        ["가명처리된 데이터", "모든 주체에게 개인정보", "상대적 접근법 — 수신자에게 비개인정보일 수 있음"],
    ])

    add_paragraph(doc,
        "다만, 디지털 옴니버스는 제안 단계에 있으며, 유럽의회와 이사회가 아직 3자 협상"
        "(trilogue)을 시작하지 아니하였습니다. 최종 조문은 상당히 달라질 수 있으므로, "
        "현재 준수 결정에 의존할 수 없음을 유의하여야 합니다.",
        left_indent=Cm(1.0))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 4. 결론 및 권고사항 (§3.5)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "4. 결론 및 권고사항", bold=True, size=Pt(13), space_before=Pt(18), space_after=Pt(8))

    add_paragraph(doc,
        "위의 사실관계와 관련 법령, EDPB 가이드라인 및 집행 사례를 종합하여 보건대, "
        "귀사는 현재 GDPR 제28조상 개인정보 수탁처리자(processor) 자격으로 고객 조달 데이터를 "
        "사용하여 LLM을 미세조정하는 데 정당한 이익(Art. 6(1)(f))을 법적 근거로 원용할 수 "
        "없을 것으로 사료됩니다. 수탁처리자가 고객 데이터를 자사의 AI 학습 목적으로 사용하는 "
        "것은 수탁처리 관계의 근본적 위반에 해당하며, 해당 처리에 대해 사실상 개인정보처리자"
        "(controller)로 간주될 것입니다.",
        left_indent=Cm(0.5))

    add_paragraph(doc,
        "가. 즉시 조치 (현행법)",
        bold=True, size=Pt(12), space_before=Pt(12), space_after=Pt(6),
        left_indent=Cm(0.5))

    recs_immediate = [
        "현재의 수탁처리자 자격으로 AI 모델 학습을 진행하지 마십시오. 이는 제28조 제10항에 따른 위법한 목적 결정에 해당합니다.",
        "고객과의 법률관계를 재구성하십시오:\n"
        "   · 방안 A: 각 고객에게 AI 모델 학습을 위한 데이터 사용에 대한 명시적 계약상 허가를 받을 것 "
        "(이 목적에 대한 공동 개인정보처리자 또는 독립적 개인정보처리자 관계를 반영하도록 DPA 수정)\n"
        "   · 방안 B: 고객을 통한 정보주체 동의 확보 — B2B 맥락에서 기술적으로 복잡하고 운영상 부담이 큼\n"
        "   · 방안 C: 데이터의 사전 익명화 — 학습 전에 조달 데이터를 완전히 익명화(모든 개인 식별정보 제거, "
        "데이터 집계)할 수 있다면 GDPR이 적용되지 아니함",
        "데이터 최소화 평가를 수행하십시오 — 개인정보가 모델 학습에 진정으로 필요한지, 또는 익명화/가명처리된 "
        "거래 데이터로 충분한지 판단하십시오.",
        "개인정보가 필요한 경우, 처리 개시 전에 전체 DPIA를 수행하십시오 (Art. 35).",
        "강력한 완화 조치를 구현하십시오:\n"
        "   · 학습 전 모든 개인 식별정보 가명처리\n"
        "   · 차등 프라이버시(Differential Privacy) 기술 적용\n"
        "   · 추출 및 멤버십 추론 공격에 대한 모델 테스트\n"
        "   · 제21조 이의제기권 메커니즘 구현\n"
        "   · 개인정보 처리방침 업데이트",
        "정당한 이익 평가를 철저히 문서화하십시오 — 이는 책임성 원칙(Art. 5(2))에 따라 요구됩니다.",
    ]
    for idx, rec in enumerate(recs_immediate, 1):
        add_paragraph(doc, f"{idx}. {rec}", left_indent=Cm(1.0))

    add_paragraph(doc,
        "나. 중기 조치 (디지털 옴니버스 대비)",
        bold=True, size=Pt(12), space_before=Pt(12), space_after=Pt(6),
        left_indent=Cm(0.5))

    recs_mid = [
        "COM(2025) 837의 입법 진행 상황을 모니터링하십시오. 제안대로 채택되면 법적 근거 분석이 "
        "용이해지나 면제되지는 아니합니다.",
        "익명화 우선 아키텍처를 고려하십시오 — 기본적으로 데이터를 익명화하는 AI 학습 파이프라인을 "
        "설계하고, 명백히 필요한 경우에만 개인정보를 사용하십시오. 이 접근법은 현행법과 제안된 법 "
        "모두에서 견고합니다.",
        "DPO와 협의하고, 제안된 AI 학습 활동에 대해 주관 감독기관(lead supervisory authority)으로부터 "
        "지침을 구하는 것을 고려하십시오.",
    ]
    for idx, rec in enumerate(recs_mid, 7):
        add_paragraph(doc, f"{idx}. {rec}", left_indent=Cm(1.0))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # §10.1 종결 Disclaimer
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "", space_before=Pt(18))  # spacing

    add_paragraph(doc,
        "본 분석 메모는 귀사의 요청 맥락을 기준으로 작성된 검토 문서이므로, 귀사 이외의 "
        "제3자가 그대로 원용하거나 전용하기에 적합하지 않을 수 있습니다. 본 문서는 별도 "
        "검토 없이 다른 목적이나 용도로 재사용되어서는 안 됩니다. 이상의 내용을 귀사의 업무에 "
        "참고하시기 바랍니다. 끝.",
        size=Pt(11))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # §10.3 서명 블록
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "", space_before=Pt(24))
    add_paragraph(doc, "Jinju Legal Orchestrator", bold=True, size=Pt(11))
    add_paragraph(doc, "", space_before=Pt(12))
    add_paragraph(doc, "EU 데이터 보호 스페셜리스트 Kim De Bruyne (김덕배)", size=Pt(11))

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # Appendix: 관련 법조문 상호참조 (new page)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_page_break()
    add_paragraph(doc, "[별첨] 관련 법조문 상호참조", bold=True, size=Pt(13), space_after=Pt(12))

    xref_data = [
        ["조항", "관련성"],
        ["GDPR Art. 4(1)", "개인정보의 정의"],
        ["GDPR Art. 4(7)", "개인정보처리자의 정의"],
        ["GDPR Art. 5(1)(a)", "적법성, 공정성, 투명성"],
        ["GDPR Art. 5(1)(b)", "목적 제한"],
        ["GDPR Art. 5(1)(c)", "데이터 최소화"],
        ["GDPR Art. 5(2)", "책임성"],
        ["GDPR Art. 6(1)(f)", "정당한 이익"],
        ["GDPR Art. 6(4)", "호환 목적 테스트"],
        ["GDPR Art. 9", "특수범주 데이터"],
        ["GDPR Art. 21", "이의제기권"],
        ["GDPR Art. 22", "자동화된 의사결정"],
        ["GDPR Art. 25", "설계에 의한 데이터 보호"],
        ["GDPR Art. 28", "수탁처리자 의무"],
        ["GDPR Art. 29", "개인정보처리자 권한 하의 처리"],
        ["GDPR Art. 35", "DPIA"],
        ["GDPR Art. 82", "손해배상 청구권"],
        ["GDPR Art. 83(4)(a)", "수탁처리자 의무 위반 과징금"],
        ["GDPR Art. 89", "연구/통계 목적의 보호조치"],
        ["GDPR Recital 39", "목적 제한, 데이터 최소화"],
        ["GDPR Recital 47", "정당한 이익 / 합리적 기대"],
        ["GDPR Recital 50", "호환 목적"],
        ["GDPR Recital 81", "수탁처리자 의무"],
        ["EU AI Act Art. 10", "고위험 AI의 데이터 거버넌스"],
        ["COM(2025) 837", "디지털 옴니버스 GDPR 개정안 (제안)"],
    ]
    add_analysis_table(doc, xref_data)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # Appendix: 인용 주요 출처
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    add_paragraph(doc, "[별첨] 인용 주요 출처", bold=True, size=Pt(13), space_before=Pt(18), space_after=Pt(8))

    add_paragraph(doc, "EDPB 문서", bold=True, left_indent=Cm(0.5), space_before=Pt(8))
    edpb_sources = [
        "EDPB Opinion 28/2024 — AI models and personal data (2024. 12. 17.)",
        "EDPB Guidelines 1/2024 — Art. 6(1)(f) legitimate interest (2024. 10. 8.)",
        "EDPB Guidelines 07/2020 — Concepts of controller and processor",
        "EDPB ChatGPT Taskforce Report (2024. 5. 23.)",
    ]
    for s in edpb_sources:
        add_paragraph(doc, f"· {s}", left_indent=Cm(1.0))

    add_paragraph(doc, "CJEU 판례", bold=True, left_indent=Cm(0.5), space_before=Pt(8))
    cjeu_sources = [
        "C-621/22 — KNLTB (legitimate interest, commercial purposes)",
        "C-252/21 — Meta v Bundeskartellamt (legitimate interest, tracking)",
        "C-683/21 — Nacionalinis visuomenės sveikatos centras (joint controller, processor liability)",
        "C-40/17 — Fashion ID (joint controller without data access)",
        "C-210/16 — Wirtschaftsakademie (joint controller, Facebook fan pages)",
        "C-634/21 — SCHUFA Scoring (Art. 22 automated credit scoring)",
        "C-413/23 — EDPS v SRB (relative approach to personal data/pseudonymisation)",
        "C-582/14 — Breyer (dynamic IP addresses, legitimate interest)",
    ]
    for s in cjeu_sources:
        add_paragraph(doc, f"· {s}", left_indent=Cm(1.0))

    add_paragraph(doc, "집행 결정", bold=True, left_indent=Cm(0.5), space_before=Pt(8))
    add_paragraph(doc, "· Garante v OpenAI — EUR 15 million (2024. 12. 20.)", left_indent=Cm(1.0))

    add_paragraph(doc, "입법 제안", bold=True, left_indent=Cm(0.5), space_before=Pt(8))
    add_paragraph(doc, "· COM(2025) 837 — Digital Legal Omnibus", left_indent=Cm(1.0))

    return doc


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    doc = build_document()
    output_path = OUTPUT_DIR / '2026-03-26-ai-model-training-legitimate-interest-KO.docx'
    doc.save(str(output_path))
    print(f'Generated: {output_path}')
